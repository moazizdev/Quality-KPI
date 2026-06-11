import io
from collections import defaultdict
from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import Optional
from datetime import date, timedelta

from app.database import get_db
from app import models
from app.auth import get_current_user
from pathlib import Path

router = APIRouter(prefix="/reports", tags=["Reports"], dependencies=[Depends(get_current_user)])


def get_week_range(week_ref: Optional[date] = None):
    if week_ref is None:
        week_ref = date.today()
    days_since_saturday = (week_ref.weekday() - 5) % 7
    week_start = week_ref - timedelta(days=days_since_saturday)
    week_end = week_start + timedelta(days=6)
    return week_start, week_end


HALL_TRANSLATIONS = {
    "صالة إنتاج 1": "صالة إنتاج 1",
    "صالة إنتاج 2": "صالة إنتاج 2",
    "صالة إنتاج 3": "صالة إنتاج 3",
}


def _build_dept_data(db: Session):
    depts = db.query(models.Department).order_by(models.Department.sort_order, models.Department.name).all()
    keywords = {}
    translations = {}
    for d in depts:
        if d.defect_prefixes:
            keywords[d.name] = [p.strip() for p in d.defect_prefixes.split(",") if p.strip()]
        if d.name_en:
            translations[d.name_en] = d.name
    return keywords, translations


def build_report_data(db: Session, week_start: Optional[date] = None) -> dict:
    ws, we = get_week_range(week_start)
    dept_keywords, dept_translations = _build_dept_data(db)

    halls = db.query(models.Hall).order_by(models.Hall.name).all()
    machines = db.query(models.Machine).order_by(models.Machine.hall_id, models.Machine.machine_code).all()
    hall_map = {h.id: HALL_TRANSLATIONS.get(h.name, h.name) for h in halls}

    prod_pieces = dict(
        db.query(
            models.ProductionRecord.machine_id,
            func.sum(models.ProductionRecord.pieces_produced),
        )
        .filter(
            models.ProductionRecord.production_date.between(ws, we),
            models.ProductionRecord.pieces_produced.isnot(None),
        )
        .group_by(models.ProductionRecord.machine_id)
        .all()
    )
    prod_record_counts = dict(
        db.query(
            models.ProductionRecord.machine_id,
            func.count(models.ProductionRecord.id),
        )
        .filter(models.ProductionRecord.production_date.between(ws, we))
        .group_by(models.ProductionRecord.machine_id)
        .all()
    )

    deviations = (
        db.query(models.Deviation)
        .filter(models.Deviation.date.between(ws, we))
        .all()
    )
    defect_categories = {d.id: d for d in db.query(models.DefectCategory).all()}

    dev_ids = [d.id for d in deviations]
    capa_rows = db.query(models.CAPACase).filter(
        models.CAPACase.deviation_id.in_(dev_ids)
    ).all() if dev_ids else []
    capa_by_dev = defaultdict(list)
    for c in capa_rows:
        capa_by_dev[c.deviation_id].append({
            "probable_cause": c.probable_cause or "",
            "assigned_department": c.assigned_department or "",
        })

    dev_agg = {}
    for d in deviations:
        caps = capa_by_dev.get(d.id, [{"probable_cause": "", "assigned_department": ""}])
        cap = caps[0] if caps else {"probable_cause": "", "assigned_department": ""}
        ar_dept = dept_translations.get(cap["assigned_department"], cap["assigned_department"])
        cause_ar = cap["probable_cause"]
        key = (d.machine_id, d.defect_category_id, cause_ar, ar_dept)
        dev_agg[key] = dev_agg.get(key, 0) + d.quantity

    machine_report = []
    for m in machines:
        total_pieces = prod_pieces.get(m.id, 0) or 0
        total_records = prod_record_counts.get(m.id, 0)

        machine_defects = []
        for (mid, cid, cause, dept), qty in dev_agg.items():
            if mid == m.id:
                cat = defect_categories.get(cid)
                machine_defects.append({
                    "defect_code": cat.defect_code if cat else "",
                    "defect_name": cat.defect_name if cat else "",
                    "quantity": qty,
                    "rate_percent": round((qty / total_pieces * 100), 3) if total_pieces > 0 else 0.0,
                    "probable_cause": cause,
                    "assigned_department": dept,
                })

        machine_defects.sort(key=lambda x: x["rate_percent"], reverse=True)
        total_defect_qty = sum(d["quantity"] for d in machine_defects)
        overall_rate = round((total_defect_qty / total_pieces * 100), 3) if total_pieces > 0 else 0.0

        machine_report.append({
            "machine_id": m.id,
            "machine_code": m.machine_code,
            "hall": hall_map.get(m.hall_id, ""),
            "hall_id": m.hall_id,
            "total_records": total_records,
            "total_pieces_produced": total_pieces,
            "total_defect_quantity": total_defect_qty,
            "overall_defect_rate": overall_rate,
            "defects": machine_defects,
        })

    dept_totals = defaultdict(int)
    for (mid, cid, cause, dept), qty in dev_agg.items():
        cat = defect_categories.get(cid)
        if cat:
            for dept_name, keywords in dept_keywords.items():
                if any(cat.defect_code.startswith(k) for k in keywords):
                    dept_totals[dept_name] += qty
                    break

    total_defects_all = sum(dept_totals.values())
    department_summary = [
        {
            "department": dept,
            "total_quantity": qty,
            "percentage": round((qty / total_defects_all * 100), 2) if total_defects_all > 0 else 0.0,
        }
        for dept, qty in sorted(dept_totals.items(), key=lambda x: x[1], reverse=True)
    ]

    cat_totals = defaultdict(int)
    for (mid, cid, cause, dept), qty in dev_agg.items():
        cat_totals[cid] += qty
    category_summary = []
    for cid, qty in cat_totals.items():
        cat = defect_categories.get(cid)
        if cat:
            category_summary.append({
                "defect_code": cat.defect_code,
                "defect_name": cat.defect_name,
                "total_quantity": qty,
            })
    category_summary.sort(key=lambda x: x["total_quantity"], reverse=True)

    total_pieces_all = sum(prod_pieces.values()) or 1

    return {
        "week_start": ws.isoformat(),
        "week_end": we.isoformat(),
        "ws": ws,
        "we": we,
        "total_production_records": sum(prod_record_counts.values()),
        "total_pieces_produced": sum(prod_pieces.values()),
        "total_defect_quantity": total_defects_all,
        "overall_defect_rate": round((total_defects_all / total_pieces_all * 100), 2),
        "machines": machine_report,
        "department_summary": department_summary,
        "category_summary": category_summary,
    }


@router.get("/weekly")
def weekly_report(
    week_start: Optional[date] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    data = build_report_data(db, week_start)
    data.pop("ws", None)
    data.pop("we", None)
    if current_user.role != "admin":
        data.pop("total_pieces_produced", None)
        for m in data.get("machines", []):
            m.pop("total_pieces_produced", None)
    return data


MONTH_NAMES = {
    1: "يناير", 2: "فبراير", 3: "مارس", 4: "أبريل", 5: "مايو", 6: "يونيو",
    7: "يوليو", 8: "أغسطس", 9: "سبتمبر", 10: "أكتوبر", 11: "نوفمبر", 12: "ديسمبر",
}


def _week_number(d: date) -> int:
    _, week, _ = d.isocalendar()
    return week


# ─── DOCX Export (Arabic) ─────────────────────────────────────────────────────

@router.get("/weekly/docx")
def weekly_report_docx(
    week_start: Optional[date] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    from docx import Document
    from docx.shared import Pt, RGBColor, Emu, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data = build_report_data(db, week_start)
    ws = data["ws"]
    we = data["we"]
    week_num = _week_number(ws)
    month_name = MONTH_NAMES.get(ws.month, "")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.left_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")

    def _set_run_font(run, size=11, bold=False, color=None, name="Arial"):
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:cs"), name)
        rFonts.set(qn("w:eastAsia"), name)

    def _add_heading(text, size=14, space_before=8, space_after=4, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
        p = doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        _set_run_font(run, size, bold=True)
        return p

    def _add_body(text, size=11, bold=False, space_after=3):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        _set_run_font(run, size, bold=bold)
        return p

    def _add_hrule():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1a73e8")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.autofit = True
        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)
        bidiVisual = OxmlElement("w:bidiVisual")
        tblPr.append(bidiVisual)
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    _set_run_font(run, size=10, bold=True, color=RGBColor(255, 255, 255))
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "1a73e8")
            shading.set(qn("w:val"), "clear")
            hdr[i]._tc.get_or_add_tcPr().append(shading)
        for idx, row_data in enumerate(rows):
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = str(val)
                for p in row[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.line_spacing = 1.0
                    for run in p.runs:
                        _set_run_font(run, size=10)
                if idx % 2 == 1:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "f0f4f8")
                    shading.set(qn("w:val"), "clear")
                    row[i]._tc.get_or_add_tcPr().append(shading)
        return table

    # ── Title ──
    _add_heading("تقرير الأداء الأسبوعى لقسم مراقبة الجودة", size=28, space_before=0, space_after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_hrule()

    # ── Period & Audience ──
    period_text = f"الفترة : {month_name} {ws.year} (الأسبوع {week_num})  ({ws} – {we})"
    _add_heading(period_text, size=13, space_before=4, space_after=1)
    _add_heading("الجهة الموجه إليها : الإدارة العليا", size=13, space_before=0, space_after=6)

    # ── Intro ──
    _add_body(
        "خلال هذا الأسبوع تم تفعيل نظام متكامل لمراقبة الجودة فى صالات الإنتاج يربط "
        "بين سجلات الفحص اليومية ونظام الإجراءات التصحيحية (CAPA).",
    )
    _add_body("الهدف هو :", bold=True)
    goals = [
        "التعرف على العيوب وأسبابها الجذرية.",
        "حلها بشكل يضمن جودة المنتج وتقليل الخسائر لأبعد مدى ممكن.",
        "ضمان مطابقة المنتج للمواصفات القياسية.",
    ]
    for g in goals:
        p = doc.add_paragraph(style="List Paragraph")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(g)
        _set_run_font(run, size=11)

    doc.add_paragraph()

    # ── KPI Summary ──
    total_pieces = data["total_pieces_produced"] or 0
    total_defects = data["total_defect_quantity"]
    overall_rate = data["overall_defect_rate"]

    kpi_text = (
        f"إجمالى القطع المنتجة : {total_pieces:,}  |  "
        f"إجمالى العيوب : {total_defects:,}  |  "
        f"نسبة العيوب الإجمالية : {overall_rate}%"
    )
    _add_body(kpi_text, bold=True)
    doc.add_paragraph()

    # ── Weight Compliance Analysis ──
    _add_heading("أولاً: تحليل أداء الأوزان (Weight Compliance Analysis):", size=13, space_after=2)
    _add_body(
        "يتم تقييم الأداء بناءً على مدى الالتزام بالمدى المسموح به (Target Range) "
        'كما هو موضح فى شيت "QC Indicators" المرفق مع التقرير.',
    )
    doc.add_paragraph()

    # ── Defect Analysis ──
    _add_heading("ثانياً: تحليل العيوب (Defect Analysis):", size=13, space_after=2)
    _add_body(
        "بناءا على سجلات الحيود تم تصنيف العيوب الأكثر تأثيرا فى صالات الإنتاج :",
    )
    doc.add_paragraph()

    # ── Group machines by hall ──
    hall_machines = defaultdict(list)
    for m in data["machines"]:
        hall_machines[m["hall"]].append(m)

    for hall_name in sorted(hall_machines.keys()):
        _add_heading(f"{hall_name} :", size=16, space_after=2)
        doc.add_paragraph()

        for m in hall_machines[hall_name]:
            _add_heading(f"ماكينة {m['machine_code']} :", size=13, space_after=2)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(
                f"القطع المنتجة: {m['total_pieces_produced']:,}  |  "
                f"إجمالى العيوب: {m['total_defect_quantity']:,}  |  "
                f"نسبة العيوب: {m['overall_defect_rate']}%"
            )
            _set_run_font(run, size=10, bold=True, color=RGBColor(100, 116, 139))

            if m["defects"]:
                _add_table(
                    ["العيب Defect", "سبب العيب", "الإدارة الموجه لها", "نسبة العيوب %"],
                    [
                        [
                            d["defect_name"],
                            d["probable_cause"],
                            d["assigned_department"],
                            f"{d['rate_percent']}%",
                        ]
                        for d in m["defects"]
                    ],
                )
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run("لا توجد عيوب مسجلة.")
                _set_run_font(run, size=10, bold=False)
                run.italic = True

            doc.add_paragraph()

        # Hall summary
        hall_total_qty = sum(m["total_defect_quantity"] for m in hall_machines[hall_name])
        hall_total_pieces = sum(m["total_pieces_produced"] for m in hall_machines[hall_name])
        hall_rate = round((hall_total_qty / hall_total_pieces * 100), 2) if hall_total_pieces > 0 else 0.0

        _add_body(
            f"يتلخص إجمالى عيوب {hall_name} والتى تم رصدها خلال الأسبوع فى أنها :  "
            f"إجمالى القطع المنتجة {hall_total_pieces:,} قطعة، إجمالى العيوب {hall_total_qty:,} "
            f"بمعدل عيب {hall_rate}%",
        )

        # Collect unique defect summaries per hall
        all_defects = []
        for m in hall_machines[hall_name]:
            for d in m["defects"]:
                all_defects.append(d)
        all_defects.sort(key=lambda x: x["rate_percent"], reverse=True)

        dept_issues = defaultdict(list)
        for d in all_defects[:10]:
            dept = d["assigned_department"] if d["assigned_department"] else "غير محدد"
            dept_issues[dept].append(d["defect_name"])

        for dept, issues in dept_issues.items():
            short_issues = "، ".join(list(dict.fromkeys(issues))[:3])
            _add_body(f"• {short_issues} ({dept}).")

        doc.add_paragraph()

    # ── Department Distribution ──
    _add_heading("ويرفق لسيادتكم نسبة العيوب الموجهة لكل إدارة خلال الأسبوع الحالى :", size=13, space_after=2)
    doc.add_paragraph()

    _add_table(
        ["الإدارة", "نسبة العيوب الموجهه لها"],
        [
            [d["department"], f"{d['percentage']}%"]
            for d in data["department_summary"]
        ],
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"weekly_report_{data['week_start']}_{data['week_end']}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── PDF Export (Arabic, RTL) ─────────────────────────────────────────────────

def _find_font(name: str) -> str:
    candidates = [
        Path(f"/usr/share/fonts/truetype/msttcorefonts/{name}"),
        Path(f"/usr/share/fonts/truetype/msttcorefonts/{name.lower()}"),
        Path(f"/usr/share/fonts/truetype/noto/{name}"),
        Path(f"/usr/share/fonts/noto/{name}"),
        Path(f"/usr/local/share/fonts/{name}"),
        Path.home() / f".fonts/{name}",
    ]
    for p in candidates:
        if p.is_file():
            return str(p)
    return ""


@router.get("/weekly/pdf")
def weekly_report_pdf(
    week_start: Optional[date] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    try:
        from weasyprint import HTML
    except ImportError:
        from fastapi.responses import JSONResponse
        return JSONResponse(
            status_code=501,
            content={"detail": "PDF generation requires weasyprint. Install it with: pip install weasyprint"},
        )

    font_regular = _find_font("Arial.ttf") or _find_font("NotoNaskhArabic-Regular.ttf")
    font_bold = _find_font("Arial_Bold.ttf") or _find_font("arialbd.ttf") or _find_font("NotoNaskhArabic-Bold.ttf")

    data = build_report_data(db, week_start)
    ws = data["ws"]
    we = data["we"]
    week_num = _week_number(ws)
    month_name = MONTH_NAMES.get(ws.month, "")

    # Group machines by hall
    hall_machines = defaultdict(list)
    for m in data["machines"]:
        hall_machines[m["hall"]].append(m)

    machine_sections = ""
    for hall_name in sorted(hall_machines.keys()):
        machine_sections += f'<h2 class="hall-title">{hall_name}</h2>'
        for m in hall_machines[hall_name]:
            if m["defects"]:
                rows = "".join(
                    f"<tr><td>{d['defect_name']}</td><td>{d['probable_cause']}</td>"
                    f"<td>{d['assigned_department']}</td><td>{d['rate_percent']}%</td></tr>"
                    for d in m["defects"]
                )
            else:
                rows = '<tr><td colspan="4" class="empty-cell">لا توجد عيوب مسجلة</td></tr>'

            machine_sections += f"""
            <div class="machine-block">
                <h3 class="machine-title">ماكينة {m['machine_code']}</h3>
                <p class="machine-summary">
                    القطع المنتجة: {m['total_pieces_produced']:,} &nbsp;|&nbsp;
                    إجمالى العيوب: {m['total_defect_quantity']:,} &nbsp;|&nbsp;
                    نسبة العيوب: {m['overall_defect_rate']}%
                </p>
                <table>
                    <thead><tr>
                        <th>العيب<br/>Defect</th>
                        <th>سبب العيب</th>
                        <th>الإدارة الموجه لها</th>
                        <th>نسبة العيوب %</th>
                    </tr></thead>
                    <tbody>{rows}</tbody>
                </table>
            </div>"""

        # Hall summary (once per hall, after all machines)
        hall_total_qty = sum(mm["total_defect_quantity"] for mm in hall_machines[hall_name])
        hall_total_pieces = sum(mm["total_pieces_produced"] for mm in hall_machines[hall_name])
        hall_rate = round((hall_total_qty / hall_total_pieces * 100), 2) if hall_total_pieces > 0 else 0.0

        all_defects = []
        for mm in hall_machines[hall_name]:
            for d in mm["defects"]:
                all_defects.append(d)
        all_defects.sort(key=lambda x: x["rate_percent"], reverse=True)

        dept_issues = defaultdict(list)
        for d in all_defects[:10]:
            dept = d["assigned_department"] if d["assigned_department"] else "غير محدد"
            dept_issues[dept].append(d["defect_name"])

        summary_items = ""
        for dept, issues in dept_issues.items():
            short_issues = "، ".join(list(dict.fromkeys(issues))[:3])
            summary_items += f"<li>{short_issues} ({dept}).</li>"

        machine_sections += f"""
        <p class="hall-summary">
            يتلخص إجمالى عيوب {hall_name} والتى تم رصدها خلال الأسبوع فى أنها :
            إجمالى القطع المنتجة {hall_total_pieces:,} قطعة، إجمالى العيوب {hall_total_qty:,}
            بمعدل عيب {hall_rate}%
        </p>
        <ul class="defect-list">{summary_items}</ul>"""

    dept_rows = "".join(
        f"<tr><td>{d['department']}</td><td>{d['percentage']}%</td></tr>"
        for d in data["department_summary"]
    )

    goals_html = """
    <ul class="goals">
        <li>التعرف على العيوب وأسبابها الجذرية.</li>
        <li>حلها بشكل يضمن جودة المنتج وتقليل الخسائر لأبعد مدى ممكن.</li>
        <li>ضمان مطابقة المنتج للمواصفات القياسية.</li>
    </ul>"""

    total_pieces = data["total_pieces_produced"] or 0
    total_defects = data["total_defect_quantity"]
    overall_rate = data["overall_defect_rate"]

    html_str = f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="utf-8">
<style>
    @page {{
        size: A4;
        margin: 2.54cm;
    }}
    @font-face {{
        font-family: 'Arial';
        src: url('{font_regular}') format('truetype');
        font-weight: normal;
    }}
    @font-face {{
        font-family: 'Arial';
        src: url('{font_bold}') format('truetype');
        font-weight: bold;
    }}
    * {{ box-sizing: border-box; }}
    body {{
        font-family: 'Arial', sans-serif;
        font-size: 11px;
        color: #1e293b;
        line-height: 1.15;
        direction: rtl;
        margin: 0;
        padding: 0;
    }}
    .title {{
        text-align: center;
        font-size: 28px;
        font-weight: bold;
        margin: 0 0 8px;
        color: #1e293b;
    }}
    .title-line {{
        border: none;
        border-top: 2px solid #1a73e8;
        margin: 0 0 10px;
    }}
    .meta {{
        font-size: 13px;
        font-weight: bold;
        margin: 4px 0 2px;
    }}
    .body-text {{
        font-size: 11px;
        font-weight: normal;
        margin: 4px 0;
        line-height: 1.5;
    }}
    .body-text-bold {{
        font-size: 11px;
        font-weight: bold;
        margin: 4px 0;
    }}
    .section-head {{
        font-size: 13px;
        font-weight: bold;
        margin: 10px 0 3px;
    }}
    .hall-title {{
        font-size: 16px;
        font-weight: bold;
        margin: 16px 0 4px;
        padding: 4px 0;
        border-bottom: 2px solid #1a73e8;
    }}
    .machine-title {{
        font-size: 13px;
        font-weight: bold;
        margin: 10px 0 2px;
    }}
    .machine-summary {{
        font-size: 10px;
        font-weight: bold;
        color: #64748b;
        margin: 0 0 6px;
    }}
    .hall-summary {{
        font-size: 11px;
        font-weight: normal;
        margin: 8px 0 4px;
    }}
    .defect-list {{
        margin: 2px 0 10px 20px;
        padding-right: 20px;
        font-size: 11px;
    }}
    .defect-list li {{
        margin: 2px 0;
    }}
    .kpi-text {{
        font-size: 11px;
        font-weight: bold;
        margin: 8px 0;
    }}
    .dept-intro {{
        font-size: 13px;
        font-weight: bold;
        margin: 14px 0 4px;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin-bottom: 8px;
        font-size: 10px;
    }}
    th {{
        background: #1a73e8;
        color: #fff;
        padding: 5px 6px;
        text-align: center;
        font-weight: bold;
        font-size: 10px;
    }}
    td {{
        padding: 4px 6px;
        border-bottom: 1px solid #e2e8f0;
        text-align: center;
        font-size: 10px;
    }}
    tr:nth-child(even) td {{
        background: #f0f4f8;
    }}
    .machine-block {{
        margin-bottom: 6px;
    }}
    .empty-cell {{
        color: #64748b;
        font-size: 10px;
    }}
    .goals {{
        margin: 4px 0 8px 20px;
        padding-right: 20px;
        font-size: 11px;
    }}
    .goals li {{
        margin: 2px 0;
    }}
</style>
</head>
<body>
    <h1 class="title">تقرير الأداء الأسبوعى لقسم مراقبة الجودة</h1>
    <hr class="title-line"/>

    <p class="meta">الفترة : {month_name} {ws.year} (الأسبوع {week_num})  ({ws} – {we})</p>
    <p class="meta">الجهة الموجه إليها : الإدارة العليا</p>

    <p class="body-text">
        خلال هذا الأسبوع تم تفعيل نظام متكامل لمراقبة الجودة فى صالات الإنتاج يربط
        بين سجلات الفحص اليومية ونظام الإجراءات التصحيحية (CAPA).
    </p>
    <p class="body-text-bold">الهدف هو :</p>
    {goals_html}

    <p class="kpi-text">
        إجمالى القطع المنتجة : {total_pieces:,}  |  إجمالى العيوب : {total_defects:,}  |  نسبة العيوب الإجمالية : {overall_rate}%
    </p>

    <p class="section-head">أولاً: تحليل أداء الأوزان (Weight Compliance Analysis):</p>
    <p class="body-text">
        يتم تقييم الأداء بناءً على مدى الالتزام بالمدى المسموح به (Target Range)
        كما هو موضح فى شيت "QC Indicators" المرفق مع التقرير.
    </p>

    <p class="section-head">ثانياً: تحليل العيوب (Defect Analysis):</p>
    <p class="body-text">
        بناءا على سجلات الحيود تم تصنيف العيوب الأكثر تأثيرا فى صالات الإنتاج :
    </p>

    {machine_sections}

    <p class="dept-intro">ويرفق لسيادتكم نسبة العيوب الموجهة لكل إدارة خلال الأسبوع الحالى :</p>
    <table>
        <thead><tr><th>الإدارة</th><th>نسبة العيوب الموجهه لها</th></tr></thead>
        <tbody>{dept_rows}</tbody>
    </table>
</body>
</html>"""

    buf = io.BytesIO()
    HTML(string=html_str).write_pdf(buf)
    buf.seek(0)

    filename = f"weekly_report_{data['week_start']}_{data['week_end']}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
